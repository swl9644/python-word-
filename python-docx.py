from docx import Document
doc = Document(r"C:\Users\Administrator\Desktop\排版\完成\马克思主义哲学全书 .docx")
# print(doc1.paragraphs[100])
# print(doc2.paragraphs[100])
paragraphs= doc.paragraphs
for paragraph in paragraphs:
    if paragraph.style.name == "Heading 1":
        with open(r"C:\Users\Administrator\Desktop\排版\原版目录.txt", "a", encoding="utf-8") as f:
            f.writelines(paragraph.text + "\n")
    elif paragraph.style.name == "Heading 2":
        with open(r"C:\Users\Administrator\Desktop\排版\原版目录.txt", "a", encoding="utf-8") as f:
            f.writelines(paragraph.text + "\n")
    elif paragraph.style.name == "Heading 3":
        with open(r"C:\Users\Administrator\Desktop\排版\原版目录.txt", "a", encoding="utf-8") as f:
            f.writelines(paragraph.text + "\n")
    elif paragraph.style.name == "Heading 4":
        with open(r"C:\Users\Administrator\Desktop\排版\原版目录.txt", "a", encoding="utf-8") as f:
            f.writelines(paragraph.text + "\n")
    elif paragraph.style.name == "Heading 5":
        with open(r"C:\Users\Administrator\Desktop\排版\原版目录.txt", "a", encoding="utf-8") as f:
            f.writelines(paragraph.text + "\n")
    elif paragraph.style.name == "Heading 6":
        with open(r"C:\Users\Administrator\Desktop\排版\原版目录.txt", "a", encoding="utf-8") as f:
            f.writelines(paragraph.text + "\n")
    elif paragraph.style.name == "Heading 7":
        with open(r"C:\Users\Administrator\Desktop\排版\原版目录.txt", "a", encoding="utf-8") as f:
            f.writelines(paragraph.text + "\n")
    elif paragraph.style.name == "Heading 8":
        with open(r"C:\Users\Administrator\Desktop\排版\原版目录.txt", "a", encoding="utf-8") as f:
            f.writelines(paragraph.text + "\n")